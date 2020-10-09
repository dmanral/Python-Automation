# Read and write word document

import docx

# New document object.
d = docx.Document('docx_files/demo.docx')

# List of paragraph objects.
print(d.paragraphs)

# One paragraph object.
print(d.paragraphs[0])

# Text of the one paragraph object.
print(d.paragraphs[0].text)
print(d.paragraphs[1].text)

# Save the second paragraph object to a variable.
p = d.paragraphs[1]

# List of run objects (Runs = whenever there is a change in style.)
# p should have 4.
print(p.runs)

# Run objects also have a text member.
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)

# Can also check it
print(p.runs[1].bold)

# Changing text.
p.runs[3].underline = True
p.runs[3].text = 'italic and underlined.'

# Saving a new docx.
d.save('docx_files/demo2.docx')

# Style member
print(p.style)

# Change style of a paragraph.
p.style = 'Title'
d.save('docx_files/demo3.docx')


# Creating a brand new docx.
d2 = docx.Document()        # Blank document.

d2.add_paragraph('Hello, this is a paragraph.')
d2.add_paragraph('This is abother paragraph.')

d2.save('docx_files/demo4.docx')

p1 = d2.paragraphs[0]
# Adding another line to existing paragraph.
p1.add_run('This is a new run.')

# Should have 2 runs
print(p1.runs)

# Setting the new run to bold.
p1.runs[1].bold = True

d2.save('docx_files/demo5.docx')

# No way of inserting paragraphs or runs in between. It will add them to end.
# For now you will have to open the document you want to edit and also create
# a new one at the same time and start copying and adding.
