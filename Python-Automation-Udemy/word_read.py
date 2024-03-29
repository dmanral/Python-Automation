# The Python-Docx third-party module can read and write .docx Word files.
# Open a Word file with docx.Document()
# Access one of the Paragraph objects from the "paragraphs" member variable, which is a list of Paragraph objects.
# Paragraph objects have a "text" member variable containing the text as a string value.
# Paragraphs are composed of "runs".  The "runs" member variable of a Paragraph object contains a list of Run objects.
# Run objects also have a "text" member variable.
# Run objects have a "bold", "italic", and "underline" member variables which can be set to True or False.
# Paragraph and run objects have a "style" member variable that can be set to one of Word's built-in styles.
# Word files can be created by calling add_paragraph() and add_run() to append text content.

import docx

# Returns all the text with new line for each paragraph (No styles).
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for paragraph in doc.paragraphs:
        fullText.append(paragraph.text)
    return '\n'.join(fullText)

print(getText('docx_files/demo.docx'))
